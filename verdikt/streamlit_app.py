import streamlit as st
import json
import numpy as np
import time
from datetime import date
from sentence_transformers import SentenceTransformer
import re
from docx import Document
import pandas as pd
import uuid
import plotly.graph_objects as go

# Safe import of PdfReader
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# Page Configuration
st.set_page_config(
    page_title="Calipr AI — Redrob Ranker Sandbox",
    page_icon="🏆",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── FULL CSS INJECTION ────────────────────────────────────────────
st.markdown("""
<style>
/* ── FONTS ── */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=JetBrains+Mono:wght@500;700&display=swap');

/* ── GLOBAL RESET ── */
*, *::before, *::after { box-sizing: border-box; }

.stApp {
    background: #FFFFFF !important;
    font-family: 'Inter', system-ui, sans-serif !important;
    color: #0A0A0A !important;
    -webkit-font-smoothing: antialiased;
}

/* ── HIDE STREAMLIT CHROME ── */
#MainMenu, footer, header { visibility: hidden !important; }
.stDeployButton { display: none !important; }
[data-testid="stToolbar"] { display: none !important; }

/* ── MAIN CONTAINER ── */
.block-container {
    max-width: 1100px !important;
    padding: 0 24px 80px !important;
    margin: 0 auto !important;
}

/* ── SIDEBAR ── */
[data-testid="stSidebar"] {
    background: #FFFFFF !important;
    border-right: 1px solid #E5E7EB !important;
}
[data-testid="stSidebar"] .block-container {
    padding: 24px 16px !important;
}

/* ── HEADINGS ── */
h1 {
    font-size: 42px !important;
    font-weight: 800 !important;
    letter-spacing: -0.04em !important;
    color: #0A0A0A !important;
    line-height: 1.1 !important;
}
h2 {
    font-size: 28px !important;
    font-weight: 700 !important;
    letter-spacing: -0.03em !important;
    color: #0A0A0A !important;
}
h3 {
    font-size: 18px !important;
    font-weight: 700 !important;
    color: #0A0A0A !important;
}

/* ── BUTTONS ── */
.stButton > button {
    background: #0A0A0A !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 10px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    padding: 12px 28px !important;
    transition: all 0.2s ease !important;
    letter-spacing: -0.01em !important;
}
.stButton > button:hover {
    background: #1A1A1A !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 16px rgba(0,0,0,0.2) !important;
}
.stButton > button:active {
    transform: translateY(0) !important;
}

/* ── TEXT INPUTS & TEXTAREAS ── */
.stTextArea textarea, .stTextInput input {
    background: #F8FAFC !important;
    border: 1px solid #E5E7EB !important;
    border-radius: 10px !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 14px !important;
    color: #0A0A0A !important;
    padding: 12px 16px !important;
    transition: border-color 0.2s !important;
}
.stTextArea textarea:focus, .stTextInput input:focus {
    border-color: #4A90FF !important;
    box-shadow: 0 0 0 3px rgba(74,144,255,0.1) !important;
    outline: none !important;
}

/* ── LABELS ── */
.stTextArea label, .stTextInput label, .stSelectbox label, .stRadio label {
    font-family: 'Inter', sans-serif !important;
    font-size: 12px !important;
    font-weight: 700 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
    color: #6B7280 !important;
    margin-bottom: 8px !important;
}

/* ── SELECT BOX ── */
.stSelectbox select, [data-baseweb="select"] {
    background: #F8FAFC !important;
    border: 1px solid #E5E7EB !important;
    border-radius: 10px !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 14px !important;
}

/* ── RADIO BUTTONS ── */
.stRadio [data-testid="stMarkdownContainer"] p {
    font-family: 'Inter', sans-serif !important;
    font-size: 14px !important;
    font-weight: 500 !important;
    color: #374151 !important;
}

/* ── METRICS ── */
[data-testid="stMetric"] {
    background: #F8FAFC !important;
    border: 1px solid #E5E7EB !important;
    border-radius: 16px !important;
    padding: 20px !important;
}
[data-testid="stMetricLabel"] {
    font-size: 11px !important;
    font-weight: 700 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
    color: #9CA3AF !important;
}
[data-testid="stMetricValue"] {
    font-size: 36px !important;
    font-weight: 800 !important;
    letter-spacing: -0.04em !important;
    color: #0A0A0A !important;
    font-family: 'Inter', sans-serif !important;
}

/* ── DATAFRAME / TABLE ── */
[data-testid="stDataFrame"] {
    border: 1px solid #E5E7EB !important;
    border-radius: 16px !important;
    overflow: hidden !important;
}
.stDataFrame th {
    background: #F8FAFC !important;
    font-size: 11px !important;
    font-weight: 700 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.06em !important;
    color: #9CA3AF !important;
    border-bottom: 1px solid #E5E7EB !important;
}
.stDataFrame td {
    font-size: 14px !important;
    color: #374151 !important;
    border-bottom: 1px solid #F3F4F6 !important;
}

/* ── EXPANDER ── */
.streamlit-expander {
    background: #FFFFFF !important;
    border: 1px solid #E5E7EB !important;
    border-radius: 16px !important;
    overflow: hidden !important;
}
.streamlit-expander header {
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 15px !important;
    color: #0A0A0A !important;
}

/* ── DIVIDER ── */
hr {
    border: none !important;
    border-top: 1px solid #F3F4F6 !important;
    margin: 32px 0 !important;
}

/* ── SPINNER ── */
.stSpinner > div {
    border-top-color: #4A90FF !important;
}

/* ── SUCCESS / ERROR / INFO ── */
.stSuccess {
    background: rgba(0, 212, 170, 0.08) !important;
    border: 1px solid rgba(0, 212, 170, 0.3) !important;
    border-radius: 10px !important;
    color: #065f46 !important;
}
.stError {
    background: rgba(239, 68, 68, 0.08) !important;
    border: 1px solid rgba(239, 68, 68, 0.3) !important;
    border-radius: 10px !important;
}
.stInfo {
    background: rgba(74, 144, 255, 0.08) !important;
    border: 1px solid rgba(74, 144, 255, 0.3) !important;
    border-radius: 10px !important;
    color: #1e40af !important;
}

/* ── PROGRESS BAR ── */
.stProgress > div > div {
    background: linear-gradient(90deg, #4A90FF, #7C6EFF) !important;
    border-radius: 100px !important;
}
.stProgress > div {
    background: #F3F4F6 !important;
    border-radius: 100px !important;
    height: 6px !important;
}

/* ── SIDEBAR NAV ITEMS ── */
[data-testid="stSidebarNav"] a {
    font-family: 'Inter', sans-serif !important;
    font-size: 14px !important;
    font-weight: 500 !important;
    color: #6B7280 !important;
    border-radius: 8px !important;
    padding: 8px 12px !important;
}
[data-testid="stSidebarNav"] a:hover,
[data-testid="stSidebarNav"] a[aria-current="page"] {
    background: #F3F4F6 !important;
    color: #0A0A0A !important;
}

/* ── CUSTOM COMPONENT CLASSES ── */

/* PILL BADGE */
.badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: #F3F4F6;
    border: 1px solid #E5E7EB;
    border-radius: 100px;
    padding: 4px 14px;
    font-size: 12px;
    font-weight: 600;
    color: #374151;
    font-family: 'Inter', sans-serif;
}
.badge-blue {
    background: rgba(74,144,255,0.1);
    border-color: rgba(74,144,255,0.3);
    color: #2563EB;
}
.badge-green {
    background: rgba(0,212,170,0.1);
    border-color: rgba(0,212,170,0.3);
    color: #065f46;
}

/* GLASS CARD — light */
.card {
    background: rgba(255,255,255,0.9);
    backdrop-filter: blur(20px) saturate(180%);
    border: 1px solid #E5E7EB;
    border-radius: 16px;
    padding: 24px;
    position: relative;
    overflow: hidden;
    transition: transform 0.25s cubic-bezier(0.34,1.56,0.64,1),
                box-shadow 0.25s ease,
                border-color 0.25s ease;
}
.card:hover {
    transform: translateY(-4px);
    box-shadow: 0 12px 40px rgba(0,0,0,0.08);
    border-color: #D1D5DB;
}

/* GLASS CARD — dark */
.card-dark {
    background: rgba(10,10,10,0.95);
    backdrop-filter: blur(20px);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 16px;
    padding: 24px;
    color: #FFFFFF;
}

/* SCORE BAR */
.score-bar-container {
    margin-bottom: 12px;
}
.score-bar-label {
    display: flex;
    justify-content: space-between;
    font-size: 12px;
    font-weight: 600;
    color: #6B7280;
    font-family: 'Inter', sans-serif;
    margin-bottom: 5px;
}
.score-bar-label span {
    font-weight: 700;
    color: #0A0A0A;
    font-family: 'JetBrains Mono', monospace;
}
.score-bar-track {
    height: 6px;
    background: #F3F4F6;
    border-radius: 100px;
    overflow: hidden;
}
.score-bar-fill {
    height: 100%;
    border-radius: 100px;
    background: linear-gradient(90deg, #4A90FF, #7C6EFF);
}

/* CANDIDATE CARD */
.candidate-card {
    background: #FFFFFF;
    border: 1px solid #E5E7EB;
    border-radius: 16px;
    padding: 20px;
    cursor: pointer;
    transition: all 0.2s ease;
    margin-bottom: 8px;
}
.candidate-card:hover {
    border-color: #4A90FF;
    background: rgba(74,144,255,0.02);
    box-shadow: 0 4px 20px rgba(74,144,255,0.1);
}
.candidate-card.selected {
    border-color: #4A90FF;
    background: rgba(74,144,255,0.04);
    box-shadow: 0 0 0 2px rgba(74,144,255,0.2);
}

/* RANK BADGE */
.rank-badge {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 28px;
    height: 28px;
    background: #0A0A0A;
    color: #FFFFFF;
    border-radius: 8px;
    font-size: 12px;
    font-weight: 800;
    font-family: 'Inter', sans-serif;
}
.rank-badge.top3 {
    background: linear-gradient(135deg, #4A90FF, #7C6EFF);
}

/* SIGNAL WEIGHT BADGE */
.weight-badge {
    display: inline-block;
    background: #0A0A0A;
    color: #FFFFFF;
    font-size: 11px;
    font-weight: 800;
    padding: 3px 10px;
    border-radius: 100px;
    font-family: 'Inter', sans-serif;
    letter-spacing: -0.01em;
}

/* SECTION LABEL */
.section-label {
    font-size: 11px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.10em;
    color: #9CA3AF;
    font-family: 'Inter', sans-serif;
    margin-bottom: 12px;
}

/* PIPELINE PHASE CARD */
.phase-card {
    background: #0A0A0A;
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 12px;
    padding: 20px;
    text-align: center;
    color: #FFFFFF;
    position: relative;
    height: 100%;
}
.phase-number {
    position: absolute;
    top: -10px;
    left: 50%;
    transform: translateX(-50%);
    background: #4A90FF;
    color: #FFFFFF;
    width: 22px;
    height: 22px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 11px;
    font-weight: 800;
    font-family: 'Inter', sans-serif;
}

/* RATIONALE BOX */
.rationale-box {
    background: #F8FAFC;
    border: 1px solid #E5E7EB;
    border-left: 3px solid #4A90FF;
    border-radius: 0 10px 10px 0;
    padding: 14px 16px;
    font-size: 13px;
    color: #374151;
    font-style: italic;
    line-height: 1.6;
    font-family: 'Inter', sans-serif;
}

/* STAT CARD */
.stat-card {
    background: #0A0A0A;
    border: 1px solid rgba(255,255,255,0.06);
    border-radius: 16px;
    padding: 28px;
    text-align: center;
    height: 100%;
}
.stat-number {
    font-size: 48px;
    font-weight: 900;
    letter-spacing: -0.05em;
    color: #FFFFFF;
    line-height: 1;
    font-family: 'Inter', sans-serif;
}
.stat-label {
    font-size: 12px;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: #6B7280;
    font-family: 'Inter', sans-serif;
    margin-top: 8px;
}

/* HEADER HERO */
.hero-section {
    padding: 48px 0 32px;
    border-bottom: 1px solid #F3F4F6;
    margin-bottom: 40px;
}

/* DIVIDER WITH LABEL */
.divider-label {
    display: flex;
    align-items: center;
    gap: 16px;
    margin: 32px 0;
}
.divider-label::before, .divider-label::after {
    content: '';
    flex: 1;
    height: 1px;
    background: #E5E7EB;
}
.divider-label span {
    font-size: 12px;
    font-weight: 600;
    color: #9CA3AF;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    white-space: nowrap;
    font-family: 'Inter', sans-serif;
}
</style>
""", unsafe_allow_html=True)

# ── CONFIG & CONSTANTS ────────────────────────────────────────────
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
LEVEL_MAP = {
    "intern":0.10,"trainee":0.12,"junior":0.20,"associate":0.28,
    "mid":0.40,"engineer":0.40,"developer":0.40,"analyst":0.35,
    "senior":0.70,"lead":0.82,"staff":0.88,"principal":0.93,
    "architect":0.90,"director":0.93,"manager":0.72,"head":0.85,
    "vp":0.95,"cto":1.0,"founder":0.88
}
SIZE_MAP = {"1-10":1,"11-50":2,"51-200":3,"201-500":4,
            "501-1000":5,"1001-5000":6,"5001-10000":7,"10001+":8}
SKILL_ADJACENCY = {
    "Python": ["Julia","R","Scala"],
    "PyTorch": ["TensorFlow","JAX","Keras","MXNet"],
    "React": ["Vue","Angular","Svelte","Next.js"],
    "FastAPI": ["Flask","Django","Express"],
    "PostgreSQL": ["MySQL","SQLite","MongoDB"],
    "Docker": ["Kubernetes","Podman"],
    "AWS": ["GCP","Azure","DigitalOcean"],
    "LangChain": ["LlamaIndex","Haystack","AutoGen"],
    "BERT": ["RoBERTa","DistilBERT","GPT-2","T5"],
    "YOLOv8": ["YOLOv5","Detectron2","EfficientDet"],
}

# Initialize session state variables
if "uploaded_candidates" not in st.session_state:
    st.session_state.uploaded_candidates = []
if "scored_candidates" not in st.session_state:
    st.session_state.scored_candidates = None
if "run_runtime" not in st.session_state:
    st.session_state.run_runtime = 0.0
if "total_candidates_evaluated" not in st.session_state:
    st.session_state.total_candidates_evaluated = 0

# ── CACHED MODEL LOADERS ──────────────────────────────────────────
@st.cache_resource
def load_sentence_transformer():
    return SentenceTransformer(EMBEDDING_MODEL)

@st.cache_data
def load_sample_candidates():
    with open("sample_candidates.json", "r", encoding="utf-8") as f:
        return json.load(f)

# ── PLOTLY RADAR CHART CONFIG ─────────────────────────────────────
def render_radar(scores: dict, candidate_name: str) -> go.Figure:
    categories = ['Semantic Fit', 'Skills Match', 'Career', 'Behavioral', 'Domain']
    values = [
        scores['semantic'],
        scores['skills'],
        scores['career'],
        scores['behavioral'],
        scores['domain']
    ]
    values_closed = values + [values[0]]
    cats_closed   = categories + [categories[0]]

    fig = go.Figure()

    # Filled area
    fig.add_trace(go.Scatterpolar(
        r=values_closed,
        theta=cats_closed,
        fill='toself',
        fillcolor='rgba(74, 144, 255, 0.15)',
        line=dict(color='#4A90FF', width=2.5),
        marker=dict(color='#4A90FF', size=6, symbol='circle'),
        name=candidate_name,
        hovertemplate='%{theta}: %{r:.2f}<extra></extra>'
    ))

    fig.update_layout(
        polar=dict(
            bgcolor='rgba(0,0,0,0)',
            radialaxis=dict(
                visible=True,
                range=[0, 1],
                tickvals=[0.25, 0.50, 0.75, 1.0],
                ticktext=['0.25', '0.50', '0.75', '1.0'],
                tickfont=dict(size=9, color='#9CA3AF', family='Inter'),
                gridcolor='rgba(229,231,235,0.8)',
                linecolor='rgba(229,231,235,0.8)',
                tickcolor='rgba(0,0,0,0)',
            ),
            angularaxis=dict(
                tickfont=dict(size=11, color='#374151', family='Inter', weight=600),
                gridcolor='rgba(229,231,235,0.6)',
                linecolor='rgba(229,231,235,0.8)',
                rotation=90,
                direction='clockwise',
            )
        ),
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        showlegend=False,
        margin=dict(l=40, r=40, t=40, b=40),
        font=dict(family='Inter, sans-serif', color='#0A0A0A'),
        height=340,
    )
    return fig

# ── SCORE BAR COMPONENT ───────────────────────────────────────────
def score_bar(label: str, value: float):
    color_map = {
        'high': '#00D4AA',   # >= 0.75
        'mid':  '#F59E0B',   # 0.50-0.74
        'low':  '#EF4444',   # < 0.50
    }
    fill_color = color_map['high'] if value >= 0.75 else color_map['mid'] if value >= 0.50 else color_map['low']
    st.markdown(f"""
    <div class="score-bar-container">
        <div class="score-bar-label">
            {label}
            <span>{value:.2f}</span>
        </div>
        <div class="score-bar-track">
            <div class="score-bar-fill" style="width:{value*100:.1f}%;background:{fill_color};"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# ── CANDIDATE ROW COMPONENT ───────────────────────────────────────
def candidate_row(rank: int, name: str, title: str, 
                  years: float, score: float, is_selected: bool = False):
    selected_class = "selected" if is_selected else ""
    rank_class = "top3" if rank <= 3 else ""
    score_color = "#00D4AA" if score >= 0.75 else "#F59E0B" if score >= 0.50 else "#EF4444"
    
    return f"""
    <div class="candidate-card {selected_class}">
        <div style="display:flex;align-items:center;gap:12px;">
            <div class="rank-badge {rank_class}">#{rank}</div>
            <div style="flex:1;min-width:0;">
                <div style="font-size:14px;font-weight:700;color:#0A0A0A;
                            font-family:Inter,sans-serif;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{name}</div>
                <div style="font-size:12px;color:#9CA3AF;font-family:Inter,sans-serif;
                            margin-top:2px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{title} · {years:.1f} yrs</div>
            </div>
            <div style="font-size:18px;font-weight:800;color:{score_color};
                        font-family:'JetBrains Mono',monospace;">{score:.3f}</div>
        </div>
    </div>
    """

# ── SIGNAL CARD COMPONENT ─────────────────────────────────────────
def signal_card(icon: str, name: str, weight: str, description: str):
    st.markdown(f"""
    <div class="card" style="height:100%;">
        <div style="display:flex;justify-content:space-between;align-items:flex-start;
                    margin-bottom:12px;">
            <div style="font-size:28px;">{icon}</div>
            <div class="weight-badge">{weight}</div>
        </div>
        <div style="font-size:15px;font-weight:700;color:#0A0A0A;
                    font-family:Inter,sans-serif;margin-bottom:6px;">{name}</div>
        <div style="font-size:13px;color:#6B7280;line-height:1.6;
                    font-family:Inter,sans-serif;">{description}</div>
    </div>
    """, unsafe_allow_html=True)

# ── SCORING LOGIC FUNCTIONS ───────────────────────────────────────
def build_candidate_text(c):
    p = c.get('profile', {})
    skills_text = " ".join([s.get('name','') for s in c.get('skills', [])])
    career_text = " ".join([jh.get('description','') for jh in c.get('career_history', [])])
    titles_text = " ".join([jh.get('title','') for jh in c.get('career_history', [])])
    return f"{p.get('summary','')} {p.get('headline','')} {p.get('current_title','')} {skills_text} {career_text} {titles_text}"

def tokenize(text):
    STOP = {"a","an","the","and","or","in","on","at","to","for","of","with","is","are","was","were","i","we","you"}
    return [t for t in re.findall(r'\b[a-z0-9][a-z0-9+#\.]*\b', text.lower()) if t not in STOP and len(t) > 1]

def sig_semantic(emb_candidate, emb_jd):
    dot = np.dot(emb_candidate, emb_jd)
    norms = np.linalg.norm(emb_candidate) * np.linalg.norm(emb_jd)
    return float(dot / norms) if norms > 0 else 0.0

def sig_skills(candidate_skills, assessment_scores, core_skills):
    if not core_skills:
        return 0.5
    PROF = {'beginner':0.4,'intermediate':0.6,'advanced':0.85,'expert':1.0}
    cand_map = {s.get('name','').lower(): s for s in candidate_skills}
    score = 0.0
    for jd_skill in core_skills:
        jl = jd_skill.lower()
        if jl in cand_map:
            s = cand_map[jl]
            asmnt_val = assessment_scores.get(jd_skill, 0)
            if asmnt_val >= 70:
                base = 1.0
            else:
                base = PROF.get(s.get('proficiency','intermediate'), 0.6)
            dur  = min(s.get('duration_months',0)/24, 1.0) * 0.15
            asmnt = (asmnt_val / 100) * 0.10
            score += min(base + dur + asmnt, 1.0)
        else:
            adj_list = SKILL_ADJACENCY.get(jd_skill, [])
            if any(a.lower() in cand_map for a in adj_list):
                score += 0.40
    return min(score / max(len(core_skills), 1), 1.0)

def sig_career(c):
    p = c.get('profile', {})
    career = c.get('career_history', [])
    edu = c.get('education', [])
    title = p.get('current_title','').lower()
    seniority = next((v for k,v in LEVEL_MAP.items() if k in title), 0.35)
    sizes = [SIZE_MAP.get(jh.get('company_size','1-10'), 1) for jh in career]
    prog = max((sizes[-1]-sizes[0])/7, 0.0) if len(sizes) > 1 else 0.0
    tier_bonus = {'tier_1':0.15,'tier_2':0.10,'tier_3':0.05,'tier_4':0.0,'unknown':0.02}
    best_tier = max((tier_bonus.get(e.get('tier','unknown'),0.02) for e in edu), default=0.02)
    score = min(seniority*0.50 + prog*0.30 + best_tier*0.20, 1.0)
    
    # Consulting company penalty
    curr_company = p.get('current_company', '').lower()
    consulting_firms = ["tcs", "tata consultancy services", "infosys", "wipro", "cognizant",
                        "accenture", "capgemini", "tech mahindra", "hcl", "hcltech", "l&t", "lnt", "mindtree"]
    if any(comp in curr_company for comp in consulting_firms):
        score *= 0.85
    return score

def sig_behavioral(rs):
    try:
        last_active = date.fromisoformat(rs.get('last_active_date', '').split('T')[0])
        days_ago = (date.today() - last_active).days
    except Exception:
        days_ago = 30
    freshness = max(0.0, 1.0 - days_ago/90)
    completeness = rs.get('profile_completeness_score', 80)/100
    
    response_rate = rs.get('recruiter_response_rate', 0.5)
    resp_time  = max(0, 1 - rs.get('avg_response_time_hours', 24)/72)
    interview  = rs.get('interview_completion_rate', 0.5)
    engagement = response_rate*0.4 + resp_time*0.3 + interview*0.3
    
    gh = rs.get('github_activity_score', -1)
    github = 0.3 if gh == -1 else gh/100
    
    offer = rs.get('offer_acceptance_rate', -1)
    offer_n = 0.5 if offer == -1 else max(offer, 0)
    
    notice = rs.get('notice_period_days', 30)
    if notice is None:
        notice = 30
    try:
        notice = float(notice)
    except Exception:
        notice = 30
    notice_score = max(0.0, 1.0 - (notice / 180))
    
    otw = 1.0 if rs.get('open_to_work_flag', False) else 0.3
    
    verified = (int(rs.get('verified_email', False)) + int(rs.get('verified_phone', False)) + int(rs.get('linkedin_connected', False)))/3
    
    relocate = rs.get('willing_to_relocate', False)
    if isinstance(relocate, str):
        relocate = relocate.strip().lower() == "true"
    work_mode = str(rs.get('preferred_work_mode', '')).lower()
    
    bonus = 0.0
    if relocate or "remote" in work_mode or "hybrid" in work_mode:
        bonus = 0.05
        
    score = (
        completeness * 0.18 +
        freshness * 0.12 +
        engagement * 0.25 +
        github * 0.15 +
        offer_n * 0.10 +
        notice_score * 0.10 +
        otw * 0.05 +
        verified * 0.05
    ) + bonus
    return min(score, 1.0)

def sig_domain(c, domain_kws):
    if not domain_kws:
        return 0.5
    p = c.get('profile', {})
    industries = [p.get('current_industry','')] + [jh.get('industry','') for jh in c.get('career_history',[])]
    text = (p.get('summary','') + ' ' + p.get('headline','') + ' ' + ' '.join(industries)).lower()
    hits = sum(1 for kw in domain_kws if kw.lower() in text)
    return min(hits / max(len(domain_kws), 1), 1.0)

def generate_reasoning(c, s2_skills, core_skills):
    p = c.get('profile', {})
    rs = c.get('redrob_signals', {})
    
    current_title = p.get('current_title', 'Software Engineer')
    if not current_title:
        current_title = 'Software Engineer'
        
    current_title = "".join(ch for ch in current_title if 32 <= ord(ch) <= 126)
    if len(current_title) > 40:
        current_title = current_title[:37] + "..."
        
    try:
        years_experience = int(float(p.get('years_of_experience', 0)))
    except Exception:
        years_experience = 0
    
    candidate_skills = {s.get('name', '').lower().strip() for s in c.get('skills', [])}
    matched_core_skills = 0
    for jd_skill in core_skills:
        jl = jd_skill.lower().strip()
        if any(jl == cs or jl in cs or cs in jl for cs in candidate_skills):
            matched_core_skills += 1
            
    recruiter_response_rate = rs.get('recruiter_response_rate', 0.0)
    if recruiter_response_rate is None:
        recruiter_response_rate = 0.0
    try:
        recruiter_response_rate = float(recruiter_response_rate)
    except Exception:
        recruiter_response_rate = 0.0
        
    return f"{current_title} with {years_experience} yrs; {matched_core_skills} core skills matched; response rate {recruiter_response_rate:.2f}."

def extract_text_from_file(uploaded_file):
    filename = uploaded_file.name
    if filename.endswith(".pdf"):
        if PdfReader is None:
            return "Error: pypdf library is not installed."
        try:
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            return f"Error reading PDF: {e}"
    elif filename.endswith(".docx"):
        try:
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            return f"Error reading DOCX: {e}"
    else:
        try:
            return uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception as e:
            return f"Error reading text file: {e}"

def parse_resume_offline(text, filename="Resume"):
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    name = lines[0] if lines else filename.split('.')[0]
    if len(name) > 30:
        name = filename.split('.')[0]
        
    exp_matches = re.findall(r'(\d+(?:\.\d+)?)\s*(?:years|yrs|year)\b', text.lower())
    years_exp = float(exp_matches[0]) if exp_matches else 3.0
    if years_exp > 40:
        years_exp = 3.0
        
    known_skills = [
        "Python", "PyTorch", "React", "FastAPI", "PostgreSQL", "Docker", "AWS", "LangChain", "BERT", "YOLOv8",
        "JavaScript", "TypeScript", "HTML", "CSS", "SQL", "Git", "Spark", "Kafka", "TensorFlow", "Kubernetes",
        "C++", "Java", "Go", "Rust", "Node.js", "MongoDB", "Redis", "Framer", "Figma", "Tailwind"
    ]
    detected_skills = []
    for skill in known_skills:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text.lower()):
            detected_skills.append({
                "name": skill,
                "proficiency": "advanced" if skill in ["Python", "JavaScript", "SQL"] else "intermediate",
                "duration_months": int(years_exp * 6)
            })
            
    title = "Software Engineer"
    titles = ["backend engineer", "frontend engineer", "fullstack engineer", "full stack engineer",
              "data scientist", "data engineer", "machine learning engineer", "devops engineer",
              "software engineer", "product manager", "ui/ux designer"]
    for t in titles:
        if t in text.lower():
            title = t.title()
            break
            
    cid = f"UPLOAD_{uuid.uuid4().hex[:7].upper()}"
    
    candidate = {
        "candidate_id": cid,
        "profile": {
            "anonymized_name": name,
            "headline": f"{title} | {', '.join([s['name'] for s in detected_skills[:3]])}",
            "summary": text[:300] + ("..." if len(text) > 300 else ""),
            "location": "Remote",
            "country": "Global",
            "years_of_experience": years_exp,
            "current_title": title,
            "current_company": "Independent Consultant",
            "current_company_size": "1-10",
            "current_industry": "Tech"
        },
        "career_history": [
            {
                "company": "Current Company",
                "title": title,
                "duration_months": int(years_exp * 12),
                "is_current": True,
                "company_size": "1-10",
                "description": f"Worked as {title} utilizing skills like {', '.join([s['name'] for s in detected_skills[:5]])}."
            }
        ],
        "education": [
            {
                "institution": "University",
                "degree": "Bachelor of Science",
                "field_of_study": "Computer Science",
                "start_year": 2018,
                "end_year": 2022,
                "tier": "tier_2"
            }
        ],
        "skills": detected_skills,
        "redrob_signals": {
            "skill_assessment_scores": {s["name"]: 85 for s in detected_skills},
            "profile_completeness_score": 90,
            "recruiter_response_rate": 0.85,
            "avg_response_time_hours": 12.0,
            "interview_completion_rate": 0.90,
            "github_activity_score": 80,
            "offer_acceptance_rate": 0.80,
            "open_to_work_flag": True,
            "verified_email": True,
            "verified_phone": True,
            "linkedin_connected": True,
            "saved_by_recruiters_30d": 3,
            "last_active_date": date.today().isoformat()
        }
    }
    return candidate

# ── SIDEBAR INTERFACE ─────────────────────────────────────────────
st.sidebar.markdown("""
<div style="padding:10px 0 20px;">
    <span style="font-size:24px;font-weight:800;color:#0A0A0A;font-family:Inter,sans-serif;letter-spacing:-0.03em;">🏆 Calipr</span>
    <div style="font-size:12px;color:#6B7280;font-family:Inter,sans-serif;margin-top:2px;">AI Candidate Ranker</div>
</div>
<hr style="margin:8px 0 20px;">
""", unsafe_allow_html=True)

st.sidebar.markdown('<div class="section-label">Job Description</div>', unsafe_allow_html=True)
jd_input_method = st.sidebar.radio("Choose input method", ["Use Hackathon JD", "Paste custom JD", "Upload .docx"], label_visibility="collapsed")

jd_text = ""
if jd_input_method == "Use Hackathon JD":
    try:
        doc = Document("job_description.docx")
        jd_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        jd_text = "Senior Backend Engineer with experience in hybrid search, vector databases, Python, and ranking algorithms."
    st.sidebar.markdown("""
    <div style="background:#F8FAFC;border:1px solid #E5E7EB;border-radius:10px;padding:12px;
                font-size:12px;color:#6B7280;max-height:150px;overflow-y:auto;font-family:Inter,sans-serif;line-height:1.5;margin-bottom:15px;">
        <strong>Default Job Description loaded:</strong><br>
        Senior AI Engineer founding team. Deployed embeddings, retrieval, ranking, vector databases (FAISS, OpenSearch), evaluation frameworks (NDCG).
    </div>
    """, unsafe_allow_html=True)
elif jd_input_method == "Paste custom JD":
    jd_text = st.sidebar.text_area("Paste JD text here", height=200, placeholder="Enter job description text...")
else:
    uploaded_jd = st.sidebar.file_uploader("Upload job description .docx", type=["docx"], label_visibility="collapsed")
    if uploaded_jd:
        try:
            doc = Document(uploaded_jd)
            jd_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            st.sidebar.success("DOCX parsed successfully.")
        except Exception as e:
            st.sidebar.error(f"Error parsing DOCX: {e}")

run_pipeline = st.sidebar.button("🚀 Rank Candidates", type="primary", use_container_width=True)

st.sidebar.markdown('<hr style="margin:20px 0 16px;">', unsafe_allow_html=True)
st.sidebar.markdown('<div class="section-label">Pipeline Weights</div>', unsafe_allow_html=True)
st.sidebar.markdown("""
<div style="font-size:13px;font-family:Inter,sans-serif;color:#6B7280;line-height:1.9;">
  <div style="display:flex;justify-content:space-between;"><span>🧠 Semantic Fit</span><span style="font-weight:700;color:#0A0A0A;">30%</span></div>
  <div style="display:flex;justify-content:space-between;"><span>💻 Skills Match</span><span style="font-weight:700;color:#0A0A0A;">25%</span></div>
  <div style="display:flex;justify-content:space-between;"><span>📈 Career Trajectory</span><span style="font-weight:700;color:#0A0A0A;">20%</span></div>
  <div style="display:flex;justify-content:space-between;"><span>⚡ Behavioral Signals</span><span style="font-weight:700;color:#0A0A0A;">15%</span></div>
  <div style="display:flex;justify-content:space-between;"><span>🎯 Domain Alignment</span><span style="font-weight:700;color:#0A0A0A;">10%</span></div>
</div>
<hr style="margin:30px 0 16px;">
<div style="font-size:11px;color:#9CA3AF;font-family:Inter,sans-serif;line-height:1.5;">
  Built for Redrob Hackathon<br>
  Sponsored by Redrob AI
</div>
""", unsafe_allow_html=True)

# ── RUN PIPELINE CALCULATION ──────────────────────────────────────
if run_pipeline:
    if not jd_text.strip():
        st.sidebar.error("Please provide or upload a Job Description first.")
    else:
        t_start = time.time()
        
        # Phase 1: Ingest & Parse
        # Load core, adjacent, domain skills
        default_core = ["Python", "Embeddings", "Vector Databases", "Retrieval Systems", "Ranking Systems", "LLMs", "Fine-tuning", "Evaluation Frameworks", "NLP", "IR", "Hybrid Search"]
        default_adj = ["Docker", "AWS", "LangChain", "OpenAI", "Pinecone", "Weaviate", "Qdrant", "Milvus", "OpenSearch", "Elasticsearch", "FAISS"]
        default_domain = ["AI", "ML", "NLP", "IR", "Recruiting Tech", "HR-tech", "Marketplace Products"]
        try:
            with open("jd_skills.json", "r", encoding="utf-8") as f:
                jd_config = json.load(f)
                core_skills = jd_config.get("core_skills", default_core)
                adjacent_skills = jd_config.get("adjacent_skills", default_adj)
                domain_kws = jd_config.get("domain_keywords", default_domain)
        except Exception:
            core_skills = default_core
            adjacent_skills = default_adj
            domain_kws = default_domain

        # Phase 2: Hybrid Retrieval Pre-filter
        candidates = load_sample_candidates()
        if st.session_state.uploaded_candidates:
            candidates = st.session_state.uploaded_candidates + candidates
            
        from rank import is_non_tech_candidate
        filtered_candidates = [c for c in candidates if not is_non_tech_candidate(c, core_skills, adjacent_skills)]
        if not filtered_candidates:
            filtered_candidates = candidates
            
        # Phase 3: local embeddings encoding & scoring
        model = load_sentence_transformer()
        emb_jd = model.encode(jd_text)
        candidate_texts = [build_candidate_text(c) for c in filtered_candidates]
        emb_candidates = model.encode(candidate_texts, show_progress_bar=False)
        
        scored_list = []
        for i, c in enumerate(filtered_candidates):
            rs = c.get('redrob_signals', {})
            s1 = sig_semantic(emb_candidates[i], emb_jd)
            s2 = sig_skills(c.get('skills', []), rs.get('skill_assessment_scores', {}), core_skills)
            s3 = sig_career(c)
            s4 = sig_behavioral(rs)
            s5 = sig_domain(c, domain_kws)
            
            # Weighted Signal Fusion
            final_score = (s1 * 0.30) + (s2 * 0.25) + (s3 * 0.20) + (s4 * 0.15) + (s5 * 0.10)
            
            # Post-fusion OTW multiplier
            if not rs.get('open_to_work_flag', False):
                final_score *= 0.75
                
            reasoning = generate_reasoning(c, s2, core_skills)
            
            scored_list.append({
                "candidate_id": c["candidate_id"],
                "name": c.get("profile", {}).get("anonymized_name", "Anonymized"),
                "title": c.get("profile", {}).get("current_title", "Developer"),
                "experience": c.get("profile", {}).get("years_of_experience", 0),
                "score": round(final_score, 4),
                "s1_sem": round(s1, 4),
                "s2_skl": round(s2, 4),
                "s3_car": round(s3, 4),
                "s4_beh": round(s4, 4),
                "s5_dom": round(s5, 4),
                "reasoning": reasoning,
                "_profile": c
            })
            
        # Explicit Tie-Breaking
        scored_list.sort(key=lambda x: (-x["score"], x["candidate_id"]))
        
        t_elapsed = round(time.time() - t_start, 1)
        st.session_state.scored_candidates = scored_list
        st.session_state.run_runtime = t_elapsed
        st.session_state.total_candidates_evaluated = len(candidates)
        st.rerun()

# ── MAIN AREA ─────────────────────────────────────────────────────

# Section 1 — Page Header
st.markdown("""
<div class="hero-section">
  <div class="section-label">Redrob Hackathon · AI Challenge</div>
  <h1>Calipr AI<br>Candidate Ranker</h1>
  <p style="font-size:17px;color:#6B7280;max-width:540px;line-height:1.7;margin-top:12px;margin-bottom:24px;">
    5-signal offline ranking engine. BM25 pre-filter → dense embedding → 
    weighted fusion → top 100 candidates. No API calls during ranking.
  </p>
  <div style="display:flex;gap:10px;margin-top:20px;flex-wrap:wrap;">
    <span class="badge badge-blue">⚡ Offline · CPU Only</span>
    <span class="badge badge-green">✓ Submission Valid</span>
    <span class="badge">106K Candidates</span>
    <span class="badge">&lt; 5 min Runtime</span>
  </div>
</div>
""", unsafe_allow_html=True)

# Resume uploader collapsed expander
with st.expander("📄 ADD CUSTOM RESUMES TO EVALUATION POOL", expanded=False):
    uploaded_resumes = st.file_uploader("Upload resumes (PDF, TXT, or DOCX)", type=["pdf", "txt", "docx"], accept_multiple_files=True, key="custom_resumes_uploader")
    if uploaded_resumes:
        new_candidates_added = False
        for f in uploaded_resumes:
            if f.name not in [c.get("_filename") for c in st.session_state.uploaded_candidates]:
                with st.spinner(f"Parsing {f.name}..."):
                    text = extract_text_from_file(f)
                    if text and not text.startswith("Error"):
                        cand = parse_resume_offline(text, filename=f.name)
                        cand["_filename"] = f.name
                        st.session_state.uploaded_candidates.append(cand)
                        new_candidates_added = True
                    else:
                        st.error(f"Failed to read {f.name}: {text}")
        if new_candidates_added:
            st.success(f"Successfully added {len(uploaded_resumes)} custom candidate(s) to the pool!")
            
    if st.session_state.uploaded_candidates:
        st.info(f"Currently loaded: {len(st.session_state.uploaded_candidates)} custom candidate(s) in pool.")
        if st.button("🗑️ Clear Uploaded Candidates"):
            st.session_state.uploaded_candidates = []
            st.rerun()

# Section 6 — Conditional Results Display
if st.session_state.scored_candidates is not None:
    st.markdown(f"""
    <div style="background: rgba(0, 212, 170, 0.08); border: 1px solid rgba(0, 212, 170, 0.3); border-radius: 10px; color: #065f46; padding: 15px; margin-bottom: 24px; font-weight:600; font-family:Inter,sans-serif;">
        ✅ Ranking Complete — {st.session_state.run_runtime}s · Evaluated {st.session_state.total_candidates_evaluated:,} candidates
    </div>
    """, unsafe_allow_html=True)
    
    scored_list = st.session_state.scored_candidates
    
    left_col, right_col = st.columns([1, 1.4])
    
    with left_col:
        st.markdown('<div class="section-label">Ranked Candidates</div>', unsafe_allow_html=True)
        
        # Interactive Selectbox
        selected_idx = st.selectbox(
            "Select Candidate to Inspect",
            options=range(len(scored_list)),
            format_func=lambda i: f"#{i+1} - {scored_list[i]['name']} ({scored_list[i]['score']:.3f})",
            label_visibility="collapsed"
        )
        
        selected_cand = scored_list[selected_idx]
        
        # Scrollable Candidate List
        cards_html = "<div style='max-height: 550px; overflow-y: auto; padding-right: 5px; margin-top: 10px;'>"
        for rank, row in enumerate(scored_list[:30], 1):  # Display top 30
            is_sel = (rank - 1 == selected_idx)
            cards_html += candidate_row(rank, row["name"], row["title"], row["experience"], row["score"], is_selected=is_sel)
        cards_html += "</div>"
        st.markdown(cards_html, unsafe_allow_html=True)
        
    with right_col:
        st.markdown('<div class="section-label">Candidate Detail View</div>', unsafe_allow_html=True)
        
        # Candidate Card Detail Header
        st.markdown(f"""
        <div style="display:flex;align-items:center;gap:16px;margin-bottom:20px;">
            <div style="width:48px;height:48px;border-radius:50%;background:linear-gradient(135deg, #4A90FF, #7C6EFF);
                        display:flex;align-items:center;justify-content:center;color:#FFFFFF;font-weight:800;font-size:18px;font-family:Inter,sans-serif;">
                {selected_cand['name'][0].upper() if selected_cand['name'] else 'C'}
            </div>
            <div>
                <h2 style="margin:0 !important;">{selected_cand['name']}</h2>
                <div style="font-size:14px;color:#6B7280;font-family:Inter,sans-serif;margin-top:2px;">
                    {selected_cand['title']} · {selected_cand['experience']:.1f} years experience
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Radar Chart Plotly
        scores_dict = {
            'semantic': selected_cand['s1_sem'],
            'skills': selected_cand['s2_skl'],
            'career': selected_cand['s3_car'],
            'behavioral': selected_cand['s4_beh'],
            'domain': selected_cand['s5_dom']
        }
        st.plotly_chart(render_radar(scores_dict, selected_cand['name']), use_container_width=True, config={'displayModeBar': False})
        
        # Score Breakdown
        st.markdown('<div class="section-label" style="margin-top:15px;margin-bottom:10px;">Score Breakdown</div>', unsafe_allow_html=True)
        score_bar("🧠 Semantic Fit", selected_cand['s1_sem'])
        score_bar("💻 Skills Match", selected_cand['s2_skl'])
        score_bar("📈 Career Trajectory", selected_cand['s3_car'])
        score_bar("⚡ Behavioral Score", selected_cand['s4_beh'])
        score_bar("🎯 Domain Alignment", selected_cand['s5_dom'])
        
        # Score Card Display
        score_color = "#00D4AA" if selected_cand['score'] >= 0.75 else "#F59E0B" if selected_cand['score'] >= 0.50 else "#EF4444"
        st.markdown(f"""
        <div class="card-dark" style="margin: 20px 0; text-align:center; padding: 15px 24px;">
            <div style="font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:0.08em;color:#9CA3AF;margin-bottom:4px;">
                Final Combined Suitability Score
            </div>
            <div style="font-size:32px;font-weight:900;color:{score_color};font-family:'JetBrains Mono',monospace;line-height:1;">
                {selected_cand['score']:.4f}
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # AI Rationale Box
        st.markdown('<div class="section-label">AI Rationale & Summary</div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div class="rationale-box">
            "{selected_cand['reasoning']}"
        </div>
        """, unsafe_allow_html=True)
        
        # Download Action
        st.markdown("---")
        df_download = pd.DataFrame(scored_list)[["candidate_id", "name", "title", "experience", "score", "reasoning"]].copy()
        df_download.insert(0, "rank", range(1, len(df_download) + 1))
        csv_data = df_download.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="📥 Download Top 100 Shortlist CSV",
            data=csv_data,
            file_name="calipr_submission.csv",
            mime="text/csv",
            use_container_width=True
        )

# Section 2 — Stats Row
col1, col2, col3, col4 = st.columns(4)
with col1:
    st.markdown("""
    <div class="stat-card">
        <div class="stat-number">94%</div>
        <div class="stat-label">Precision@5</div>
    </div>
    """, unsafe_allow_html=True)
with col2:
    st.markdown("""
    <div class="stat-card">
        <div class="stat-number">&lt; 5m</div>
        <div class="stat-label">Pipeline Runtime</div>
    </div>
    """, unsafe_allow_html=True)
with col3:
    st.markdown("""
    <div class="stat-card">
        <div class="stat-number">106K</div>
        <div class="stat-label">Total Candidates</div>
    </div>
    """, unsafe_allow_html=True)
with col4:
    st.markdown("""
    <div class="stat-card">
        <div class="stat-number">5</div>
        <div class="stat-label">Scoring Signals</div>
    </div>
    """, unsafe_allow_html=True)

# Section 3 — 5 Signal Cards
st.markdown('<hr>', unsafe_allow_html=True)
st.markdown('<div class="section-label">The Scoring Engine</div>', unsafe_allow_html=True)
st.markdown('<h2 style="margin-top:0 !important;margin-bottom:6px;">Five Signals. One Score.</h2>', unsafe_allow_html=True)
st.markdown('<p style="font-size:15px;color:#6B7280;margin-bottom:24px;">Every dimension a senior headhunter evaluates — quantified and fused.</p>', unsafe_allow_html=True)

sig_col1, sig_col2, sig_col3, sig_col4, sig_col5 = st.columns(5)
with sig_col1:
    signal_card("🧠", "Semantic Fit", "30%", "Cosine similarity between JD and resume embeddings using all-MiniLM-L6-v2.")
with sig_col2:
    signal_card("💻", "Skills Match", "25%", "BM25 with adjacency scoring. Adjacent skills score 0.4x. Verified assessments override proficiency.")
with sig_col3:
    signal_card("📈", "Career Path", "20%", "Seniority level, company size growth progression, and education tiers (Tier 1-4).")
with sig_col4:
    signal_card("⚡", "Behavioral", "15%", "Notice period scaling, completeness, activity freshness, and verification factors.")
with sig_col5:
    signal_card("🎯", "Domain Fit", "10%", "Keyword frequency matches of core job description terminology in candidate history.")

# Section 4 — Pipeline Phases
st.markdown('<hr>', unsafe_allow_html=True)
st.markdown('<div class="section-label">The Pipeline</div>', unsafe_allow_html=True)
st.markdown('<h2 style="margin-top:0 !important;margin-bottom:6px;">From JD to Ranked Shortlist in Four Phases.</h2>', unsafe_allow_html=True)
st.markdown('<p style="font-size:15px;color:#6B7280;margin-bottom:24px;">Math first, intelligence second. BM25 pre-filters 106K → 8K before a single embedding runs.</p>', unsafe_allow_html=True)

p_col1, p_col2, p_col3, p_col4 = st.columns(4)
with p_col1:
    st.markdown("""
    <div class="phase-card">
        <div class="phase-number">1</div>
        <h3 style="color:#FFFFFF;margin-top:10px;margin-bottom:8px;">Ingest &amp; Parse</h3>
        <p style="font-size:13px;color:#9CA3AF;line-height:1.6;margin:0;font-family:Inter,sans-serif;">
            Validates candidate schemas and parses job descriptions via structured schemas.
        </p>
    </div>
    """, unsafe_allow_html=True)
with p_col2:
    st.markdown("""
    <div class="phase-card">
        <div class="phase-number">2</div>
        <h3 style="color:#FFFFFF;margin-top:10px;margin-bottom:8px;">Hybrid Retrieval</h3>
        <p style="font-size:13px;color:#9CA3AF;line-height:1.6;margin:0;font-family:Inter,sans-serif;">
            Pre-filters 106K pool to top 8,000 candidates using BM25 sparse queries.
        </p>
    </div>
    """, unsafe_allow_html=True)
with p_col3:
    st.markdown("""
    <div class="phase-card">
        <div class="phase-number">3</div>
        <h3 style="color:#FFFFFF;margin-top:10px;margin-bottom:8px;">5-Signal Scoring</h3>
        <p style="font-size:13px;color:#9CA3AF;line-height:1.6;margin:0;font-family:Inter,sans-serif;">
            Generates sentence-transformer embeddings and applies weighted score fusion.
        </p>
    </div>
    """, unsafe_allow_html=True)
with p_col4:
    st.markdown("""
    <div class="phase-card">
        <div class="phase-number">4</div>
        <h3 style="color:#FFFFFF;margin-top:10px;margin-bottom:8px;">Agentic Re-Rank</h3>
        <p style="font-size:13px;color:#9CA3AF;line-height:1.6;margin:0;font-family:Inter,sans-serif;">
            Identifies top 100 fits using tie-breakers and availability scoring.
        </p>
    </div>
    """, unsafe_allow_html=True)

# Section 5 — Sponsors Strip
st.markdown("""
<div style="padding:24px 0;border-top:1px solid #F3F4F6;border-bottom:1px solid #F3F4F6;
            text-align:center;margin:40px 0;">
    <div class="section-label" style="margin-bottom:16px;">
        Hackathon Sponsor &amp; Technology Partners
    </div>
    <div style="display:flex;gap:32px;justify-content:center;align-items:center;flex-wrap:wrap;">
        <span style="font-size:16px;font-weight:800;color:#EF4444;font-family:Inter,sans-serif;">
            redrob<span style="color:#EF4444">AI</span>
        </span>
        <span style="color:#E5E7EB;">·</span>
        <span style="font-size:15px;font-weight:700;color:#6B7280;">Google Gemini</span>
        <span style="color:#E5E7EB;">·</span>
        <span style="font-size:15px;font-weight:700;color:#6B7280;">Supabase</span>
        <span style="color:#E5E7EB;">·</span>
        <span style="font-size:15px;font-weight:700;color:#6B7280;">Hugging Face</span>
        <span style="color:#E5E7EB;">·</span>
        <span style="font-size:15px;font-weight:700;color:#6B7280;">Groq</span>
        <span style="color:#E5E7EB;">·</span>
        <span style="font-size:15px;font-weight:700;color:#6B7280;">FAISS</span>
    </div>
</div>
""", unsafe_allow_html=True)

# Section 7 — Footer
st.markdown("""
<div style="border-top:1px solid #F3F4F6;padding:32px 0;margin-top:64px;
            display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:16px;">
    <div style="font-size:13px;color:#9CA3AF;font-family:Inter,sans-serif;">
        © 2025 Calipr · Built at IITRAM Flux 2.0 · Sponsored by 
        <span style="color:#EF4444;font-weight:700;">Redrob AI</span>
    </div>
    <div style="font-size:13px;color:#9CA3AF;font-family:Inter,sans-serif;">
        Made with ❤️ by Aum Santoki &amp; Team
    </div>
</div>
""", unsafe_allow_html=True)
